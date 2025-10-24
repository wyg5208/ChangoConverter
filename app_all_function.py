import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import os
import sys
import threading
from pathlib import Path

try:
    import pypandoc
    PANDOC_AVAILABLE = True
except ImportError:
    PANDOC_AVAILABLE = False
    pypandoc = None

# 导入版本信息
try:
    from version import VERSION
except ImportError:
    VERSION = "1.4.0"  # 备用版本号

class ChangoConverter:
    # 软件版本号（从version.py导入）
    VERSION = VERSION
    
    # Pandoc支持的所有输入格式
    INPUT_FORMATS = {
        'commonmark': 'CommonMark',
        'creole': 'Creole',
        'csv': 'CSV',
        'docbook': 'DocBook',
        'docx': 'Microsoft Word (DOCX)',
        'dokuwiki': 'DokuWiki',
        'epub': 'EPUB',
        'fb2': 'FictionBook2',
        'gfm': 'GitHub Flavored Markdown',
        'haddock': 'Haddock',
        'html': 'HTML',
        'ipynb': 'Jupyter Notebook',
        'jats': 'JATS',
        'json': 'Pandoc JSON',
        'latex': 'LaTeX',
        'markdown': 'Pandoc Markdown',
        'markdown_mmd': 'MultiMarkdown',
        'markdown_phpextra': 'PHP Markdown Extra',
        'markdown_strict': 'Strict Markdown',
        'mediawiki': 'MediaWiki',
        'man': 'Unix Man Page',
        'muse': 'Muse',
        'native': 'Pandoc Native',
        'odt': 'OpenDocument Text',
        'opml': 'OPML',
        'org': 'Emacs Org Mode',
        'rst': 'reStructuredText',
        'rtf': 'Rich Text Format',
        't2t': 'txt2tags',
        'textile': 'Textile',
        'tikiwiki': 'TikiWiki',
        'twiki': 'TWiki',
        'vimwiki': 'Vimwiki'
    }
    
    # Pandoc支持的所有输出格式
    OUTPUT_FORMATS = {
        'asciidoc': 'AsciiDoc',
        'beamer': 'LaTeX Beamer (Slides)',
        'commonmark': 'CommonMark',
        'context': 'ConTeXt',
        'docbook': 'DocBook 4',
        'docbook5': 'DocBook 5',
        'docx': 'Microsoft Word (DOCX)',
        'dokuwiki': 'DokuWiki',
        'dzslides': 'DZSlides (HTML Slides)',
        'epub': 'EPUB 2',
        'epub3': 'EPUB 3',
        'fb2': 'FictionBook2',
        'gfm': 'GitHub Flavored Markdown',
        'haddock': 'Haddock',
        'html': 'HTML 4',
        'html5': 'HTML 5',
        'icml': 'InDesign ICML',
        'ipynb': 'Jupyter Notebook',
        'jats': 'JATS',
        'json': 'Pandoc JSON',
        'latex': 'LaTeX',
        'man': 'Unix Man Page',
        'markdown': 'Pandoc Markdown',
        'markdown_mmd': 'MultiMarkdown',
        'markdown_phpextra': 'PHP Markdown Extra',
        'markdown_strict': 'Strict Markdown',
        'mediawiki': 'MediaWiki',
        'ms': 'Groff MS',
        'muse': 'Muse',
        'native': 'Pandoc Native',
        'odt': 'OpenDocument Text',
        'opml': 'OPML',
        'org': 'Emacs Org Mode',
        'pdf': 'PDF (需要LaTeX)',
        'plain': 'Plain Text',
        'pptx': 'PowerPoint',
        'revealjs': 'reveal.js (HTML Slides)',
        'rst': 'reStructuredText',
        'rtf': 'Rich Text Format',
        's5': 'S5 (HTML Slides)',
        'slideous': 'Slideous (HTML Slides)',
        'slidy': 'Slidy (HTML Slides)',
        'tei': 'TEI Simple',
        'texinfo': 'Texinfo',
        'textile': 'Textile',
        'zimwiki': 'ZimWiki'
    }
    
    # 文件扩展名到格式的映射
    EXT_TO_FORMAT = {
        '.md': 'markdown',
        '.markdown': 'markdown',
        '.mdown': 'markdown',
        '.mkd': 'markdown',
        '.docx': 'docx',
        '.html': 'html',
        '.htm': 'html',
        '.epub': 'epub',
        '.tex': 'latex',
        '.ltx': 'latex',
        '.rst': 'rst',
        '.rest': 'rst',
        '.org': 'org',
        '.odt': 'odt',
        '.rtf': 'rtf',
        '.txt': 'plain',
        '.pdf': 'pdf',
        '.pptx': 'pptx',
        '.ipynb': 'ipynb',
        '.xml': 'docbook'
    }
    
    def __init__(self, root):
        self.root = root
        self.root.title("ChangoConverter 全功能文档转换器")
        self.root.geometry("1080x480")  # 宽度1080，高度再减少20%（600→480）
        self.root.resizable(True, True)
        
        # 设置窗口图标
        try:
            icon_path = os.path.join(os.path.dirname(__file__), 'resources', 'icon.ico')
            if os.path.exists(icon_path):
                self.root.iconbitmap(icon_path)
        except Exception as e:
            pass  # 图标加载失败不影响程序运行
        
        # 输出格式选择状态（必须在create_widgets之前初始化）
        self.output_format_vars = {}
        
        # 文件覆盖策略
        self.file_overwrite_policy = None  # None=每次询问, 'overwrite'=覆盖, 'rename'=重命名, 'skip'=跳过
        self.remember_choice = False  # 是否记住用户的选择
        
        # 批量转换输入格式
        self.batch_input_format = None  # 批量转换时的输入格式选择
        
        # 初始化版本信息
        self.pandoc_version = "未安装"
        self.pandoc_installed = False
        self.latex_installed = False
        
        # 检查 Pandoc 是否可用
        self.check_pandoc_availability()
        
        # 检查 LaTeX 是否可用
        self.check_latex_availability()
        
        # 创建界面
        self.create_widgets()
        
    def check_pandoc_availability(self):
        """检查 Pandoc 是否可用"""
        if not PANDOC_AVAILABLE:
            self.pandoc_installed = False
            self.pandoc_version = "pypandoc 未安装"
            return
            
        try:
            version = pypandoc.get_pandoc_version()
            self.pandoc_version = version
            self.pandoc_installed = True
        except OSError:
            self.pandoc_installed = False
            self.pandoc_version = "未安装"
    
    def check_latex_availability(self):
        """检查 LaTeX 是否可用"""
        import subprocess
        try:
            result = subprocess.run(
                ['pdflatex', '--version'],
                capture_output=True,
                text=True,
                timeout=5
            )
            if result.returncode == 0:
                self.latex_installed = True
            else:
                self.latex_installed = False
        except (FileNotFoundError, subprocess.TimeoutExpired):
            self.latex_installed = False
            
    def create_widgets(self):
        # 主框架（减小padding）
        main_frame = ttk.Frame(self.root, padding="8")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # 配置网格权重 - 左右布局
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        
        # 创建左右两个主区域
        # 左侧区域（占2/3宽度）
        left_frame = ttk.Frame(main_frame)
        left_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(0, 5))
        left_frame.columnconfigure(1, weight=2)
        left_frame.rowconfigure(6, weight=1)
        
        # 右侧日志区域（占1/3宽度）
        right_frame = ttk.Frame(main_frame)
        right_frame.grid(row=0, column=1, sticky=(tk.W, tk.E, tk.N, tk.S))
        right_frame.columnconfigure(0, weight=1)
        right_frame.rowconfigure(1, weight=1)
        
        # 配置主框架的列权重（左2右1）
        main_frame.columnconfigure(0, weight=2)
        main_frame.columnconfigure(1, weight=1)
        main_frame.rowconfigure(0, weight=1)
        
        # 标题栏框架（放在left_frame）
        header_frame = ttk.Frame(left_frame)
        header_frame.grid(row=0, column=0, columnspan=3, pady=(0, 5))
        
        # 标题
        title_label = ttk.Label(
            header_frame, 
            text="ChangoConverter 全功能文档转换器", 
            font=('Arial', 14, 'bold')
        )
        title_label.pack(side=tk.LEFT)
        
        # 版本号标签
        version_badge = ttk.Label(
            header_frame,
            text=f"v{self.VERSION}",
            font=('Arial', 8, 'bold'),
            foreground='white',
            background='#2E86AB',
            padding=(5, 2)
        )
        version_badge.pack(side=tk.LEFT, padx=10)
        
        # 帮助图标按钮
        help_button = ttk.Button(
            header_frame,
            text="❓ 使用说明",
            command=self.show_help,
            width=10
        )
        help_button.pack(side=tk.LEFT, padx=5)
        
        # 环境检测状态框架
        status_frame = ttk.Frame(left_frame)
        status_frame.grid(row=1, column=0, columnspan=3, pady=(0, 10))
        
        # Pandoc状态
        pandoc_status_text = f"Pandoc: {self.pandoc_version}"
        pandoc_status_color = 'green' if self.pandoc_installed else 'red'
        pandoc_status_icon = '✓' if self.pandoc_installed else '✗'
        
        self.pandoc_status_label = ttk.Label(
            status_frame,
            text=f"{pandoc_status_icon} {pandoc_status_text}",
            font=('Arial', 9),
            foreground=pandoc_status_color
        )
        self.pandoc_status_label.pack(side=tk.LEFT, padx=(0, 20))
        
        # LaTeX状态
        latex_status_text = "LaTeX: 已安装" if self.latex_installed else "LaTeX: 未安装"
        latex_status_color = 'green' if self.latex_installed else 'orange'
        latex_status_icon = '✓' if self.latex_installed else '✗'
        
        self.latex_status_label = ttk.Label(
            status_frame,
            text=f"{latex_status_icon} {latex_status_text}",
            font=('Arial', 9),
            foreground=latex_status_color
        )
        self.latex_status_label.pack(side=tk.LEFT, padx=(0, 20))
        
        # 重新检测按钮
        recheck_button = ttk.Button(
            status_frame,
            text="🔄 重新检测",
            command=self.recheck_environment,
            width=12
        )
        recheck_button.pack(side=tk.LEFT, padx=5)
        
        # 开始转换按钮（移到顶部）
        self.convert_button = ttk.Button(
            status_frame,
            text="▶ 开始转换",
            command=self.start_conversion,
            style='Convert.TButton',
            width=15
        )
        self.convert_button.pack(side=tk.LEFT, padx=15)
        
        # 如果Pandoc未安装，禁用转换按钮
        if not self.pandoc_installed:
            self.convert_button.config(state='disabled')
        
        # 转换模式选择
        mode_frame = ttk.Frame(left_frame)
        mode_frame.grid(row=2, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=5)
        
        ttk.Label(mode_frame, text="转换模式:", font=('Arial', 9, 'bold')).pack(side=tk.LEFT, padx=(0, 10))
        
        self.conversion_mode = tk.StringVar(value="single")
        ttk.Radiobutton(
            mode_frame,
            text="单文件转换",
            variable=self.conversion_mode,
            value="single",
            command=self.on_mode_changed
        ).pack(side=tk.LEFT, padx=(0, 20))
        
        ttk.Radiobutton(
            mode_frame,
            text="批量转换（文件夹）",
            variable=self.conversion_mode,
            value="batch",
            command=self.on_mode_changed
        ).pack(side=tk.LEFT)
        
        # 输入文件/文件夹选择
        self.input_label = ttk.Label(left_frame, text="输入文件:")
        self.input_label.grid(row=3, column=0, sticky=tk.W, pady=5)
        self.input_path_var = tk.StringVar()
        self.input_path_var.trace('w', self.on_input_file_changed)
        input_entry = ttk.Entry(left_frame, textvariable=self.input_path_var, width=50)  # 减小宽度
        input_entry.grid(row=3, column=1, sticky=(tk.W, tk.E), padx=(10, 5), pady=5)
        self.browse_button = ttk.Button(left_frame, text="浏览", command=self.browse_input_file)
        self.browse_button.grid(row=3, column=2, padx=(0, 5), pady=5)
        
        # 自动识别的输入格式显示 + 批量转换输入格式选择
        format_info_frame = ttk.Frame(left_frame)
        format_info_frame.grid(row=4, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=5)
        
        ttk.Label(format_info_frame, text="识别格式:").pack(side=tk.LEFT, padx=(0, 5))
        self.detected_format_var = tk.StringVar()
        self.detected_format_var.set("未选择文件")
        self.format_label = ttk.Label(
            format_info_frame, 
            textvariable=self.detected_format_var,
            foreground='blue',
            font=('Arial', 9, 'bold')
        )
        self.format_label.pack(side=tk.LEFT, padx=(0, 20))
        
        # 批量转换输入格式选择（仅批量模式显示）
        self.batch_format_frame = ttk.Frame(format_info_frame)
        # 初始不显示（默认是单文件模式）
        
        ttk.Label(self.batch_format_frame, text="输入格式:", font=('Arial', 9, 'bold')).pack(side=tk.LEFT, padx=(0, 10))
        
        # 输入格式单选框
        self.batch_input_format = tk.StringVar(value="md")
        batch_formats = [
            ("Markdown", "md"),
            ("HTML", "html"),
            ("Word", "docx"),
            ("Text", "txt"),
            ("PowerPoint", "pptx"),
            ("PDF", "pdf")
        ]
        
        for text, value in batch_formats:
            ttk.Radiobutton(
                self.batch_format_frame,
                text=text,
                variable=self.batch_input_format,
                value=value
            ).pack(side=tk.LEFT, padx=(0, 10))
        
        # 输出目录选择
        self.output_dir_label = ttk.Label(left_frame, text="输出目录:")
        self.output_dir_label.grid(row=5, column=0, sticky=tk.W, pady=5)
        self.output_dir_var = tk.StringVar()
        output_dir_entry = ttk.Entry(left_frame, textvariable=self.output_dir_var, width=50)  # 减小宽度
        output_dir_entry.grid(row=5, column=1, sticky=(tk.W, tk.E), padx=(10, 5), pady=5)
        ttk.Button(left_frame, text="浏览", command=self.browse_output_dir).grid(row=5, column=2, padx=(0, 5), pady=5)
        
        # 输出目录说明（仅批量模式显示）
        self.output_dir_hint = ttk.Label(
            left_frame,
            text="💡 批量转换时：转换后的文件将保存到源文件所在文件夹",
            font=('Arial', 8),
            foreground='gray'
        )
        # 初始不显示（默认是单文件模式）
        
        # 输出格式选择区域 - 使用Notebook(TAB页面)
        output_frame = ttk.LabelFrame(left_frame, text="选择输出格式（可多选）", padding="8")
        output_frame.grid(row=6, column=0, columnspan=3, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(8, 8))
        output_frame.columnconfigure(0, weight=1)
        output_frame.rowconfigure(0, weight=1)
        
        # 创建Notebook（TAB控件）
        format_notebook = ttk.Notebook(output_frame)
        format_notebook.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 10))
        
        # 常用格式TAB
        common_tab = ttk.Frame(format_notebook)
        format_notebook.add(common_tab, text="📌 常用格式")
        
        # 常用格式列表
        common_formats_list = [
            ('markdown', 'Pandoc Markdown'),
            ('html5', 'HTML 5'),
            ('docx', 'Microsoft Word (DOCX)'),
            ('pdf', 'PDF (需要LaTeX)'),
            ('plain', 'Plain Text'),
            ('pptx', 'PowerPoint'),
            ('epub3', 'EPUB 3'),
            ('latex', 'LaTeX'),
            ('rst', 'reStructuredText'),
            ('mediawiki', 'MediaWiki'),
            ('org', 'Emacs Org Mode'),
            ('odt', 'OpenDocument Text')
        ]
        
        # 创建常用格式复选框（分4列）
        row = 0
        col = 0
        for fmt, description in common_formats_list:
            var = tk.BooleanVar(value=False)
            self.output_format_vars[fmt] = var
            
            cb_frame = ttk.Frame(common_tab)
            cb_frame.grid(row=row, column=col, sticky=(tk.W, tk.E), padx=5, pady=5)
            
            cb = ttk.Checkbutton(
                cb_frame,
                text=f"{description}",
                variable=var
            )
            cb.pack(side=tk.LEFT, fill=tk.X, expand=True)
            
            # 如果是PDF格式且LaTeX未安装，禁用该选项
            if fmt == 'pdf' and not self.latex_installed:
                cb.config(state='disabled')
                hint_label = ttk.Label(
                    cb_frame,
                    text="(需要LaTeX)",
                    font=('Arial', 8),
                    foreground='gray'
                )
                hint_label.pack(side=tk.LEFT, padx=(5, 0))
            
            col += 1
            if col >= 4:
                col = 0
                row += 1
        
        # 配置列权重
        for i in range(4):
            common_tab.columnconfigure(i, weight=1, uniform="cols")
        
        # 其他格式TAB
        other_tab = ttk.Frame(format_notebook)
        format_notebook.add(other_tab, text="📋 其他格式")
        
        # 创建滚动区域用于其他格式
        other_canvas = tk.Canvas(other_tab, height=100)
        other_scrollbar = ttk.Scrollbar(other_tab, orient="vertical", command=other_canvas.yview)
        other_scrollable = ttk.Frame(other_canvas)
        
        other_scrollable.bind(
            "<Configure>",
            lambda e: other_canvas.configure(scrollregion=other_canvas.bbox("all"))
        )
        
        other_canvas.create_window((0, 0), window=other_scrollable, anchor="nw")
        other_canvas.configure(yscrollcommand=other_scrollbar.set)
        
        # 获取已添加的常用格式的键
        common_format_keys = {fmt for fmt, _ in common_formats_list}
        
        # 添加其他格式（分4列）
        row = 0
        col = 0
        for fmt, description in sorted(self.OUTPUT_FORMATS.items()):
            if fmt in common_format_keys:
                continue
            
            var = tk.BooleanVar(value=False)
            self.output_format_vars[fmt] = var
            
            cb = ttk.Checkbutton(
                other_scrollable,
                text=f"{description}",
                variable=var
            )
            cb.grid(row=row, column=col, sticky=(tk.W, tk.E), padx=5, pady=3)
            
            col += 1
            if col >= 4:
                col = 0
                row += 1
        
        # 配置列权重
        for i in range(4):
            other_scrollable.columnconfigure(i, weight=1, uniform="cols")
        
        other_canvas.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        other_scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S))
        other_tab.rowconfigure(0, weight=1)
        other_tab.columnconfigure(0, weight=1)
        
        # 快速选择按钮（合并高级选项按钮到同一行）
        quick_select_frame = ttk.Frame(output_frame)
        quick_select_frame.grid(row=1, column=0, pady=(0, 5))
        
        # 高级选项按钮（放在左侧）
        ttk.Button(
            quick_select_frame,
            text="⚙️ 高级选项",
            command=self.show_advanced_options,
            width=12
        ).pack(side=tk.LEFT, padx=(0, 20))
        
        # 快速选择按钮（放在右侧）
        ttk.Button(quick_select_frame, text="常用格式", command=self.select_common_formats).pack(side=tk.LEFT, padx=5)
        ttk.Button(quick_select_frame, text="全选", command=self.select_all_formats).pack(side=tk.LEFT, padx=5)
        ttk.Button(quick_select_frame, text="清除", command=self.clear_all_formats).pack(side=tk.LEFT, padx=5)
        
        # 初始化高级选项变量
        self.standalone_var = tk.BooleanVar(value=True)
        self.toc_var = tk.BooleanVar(value=False)
        self.number_sections_var = tk.BooleanVar(value=False)
        self.docx_font_var = tk.StringVar(value="Times New Roman")
        self.docx_fontsize_var = tk.StringVar(value="12")
        
        # 进度条和状态（放在左侧底部）
        self.progress = ttk.Progressbar(left_frame, mode='indeterminate')
        self.progress.grid(row=7, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(5, 3))
        
        # 状态标签
        self.status_var = tk.StringVar()
        if self.pandoc_installed:
            self.status_var.set("✓ 准备就绪 - 可以开始转换")
            status_color = "green"
        else:
            self.status_var.set("✗ Pandoc 未安装 - 请先安装 Pandoc")
            status_color = "red"
        
        status_label = ttk.Label(left_frame, textvariable=self.status_var, foreground=status_color)
        status_label.grid(row=8, column=0, columnspan=3, pady=(0, 3))
        self.status_label = status_label  # 保存引用以便更新颜色
        
        # === 右侧日志区域 ===
        # 日志标题
        ttk.Label(right_frame, text="转换日志:", font=('Arial', 10, 'bold')).grid(row=0, column=0, sticky=tk.W, pady=(0, 5))
        
        # 日志文本框
        log_container = ttk.Frame(right_frame)
        log_container.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        log_container.columnconfigure(0, weight=1)
        log_container.rowconfigure(0, weight=1)
        
        self.log_text = tk.Text(log_container, wrap=tk.WORD, font=('Consolas', 9))  # 自适应高度
        log_scrollbar = ttk.Scrollbar(log_container, orient=tk.VERTICAL, command=self.log_text.yview)
        self.log_text.configure(yscrollcommand=log_scrollbar.set)
        
        self.log_text.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        log_scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S))
        
        # 添加样式
        style = ttk.Style()
        style.configure('Accent.TButton', font=('Arial', 11, 'bold'))
        
        # 转换按钮特殊样式（深绿色背景+深色文字）
        try:
            # 尝试配置醒目的按钮样式
            style.configure('Convert.TButton', 
                          font=('Arial', 12, 'bold'),
                          foreground='#1a1a1a',  # 深色文字（几乎黑色）
                          background='#006400',  # 深绿色背景（DarkGreen）
                          padding=(10, 5))
            # 某些主题可能不支持background，使用map来设置
            style.map('Convert.TButton',
                     background=[('active', '#004d00'), ('disabled', '#6c757d')],  # 更深的绿色激活状态
                     foreground=[('active', '#000000'), ('disabled', '#adb5bd')])  # 激活时黑色文字
        except:
            pass
        
        # 常用格式复选框样式（加粗字体）
        try:
            style.configure('Common.TCheckbutton', font=('Arial', 9, 'bold'))
        except:
            pass  # 某些主题可能不支持自定义复选框样式
    
    def on_mode_changed(self):
        """转换模式改变时的处理"""
        mode = self.conversion_mode.get()
        
        if mode == "single":
            # 单文件模式
            self.input_label.config(text="输入文件:")
            self.output_dir_label.grid()  # 显示输出目录
            self.output_dir_hint.grid_remove()  # 隐藏提示
            self.browse_button.config(command=self.browse_input_file)
            self.detected_format_var.set("未选择文件")
            self.format_label.pack(side=tk.LEFT, padx=(0, 20))  # 显示识别格式
            self.batch_format_frame.pack_forget()  # 隐藏批量格式选择
        else:
            # 批量转换模式
            self.input_label.config(text="输入文件夹:")
            self.output_dir_label.grid_remove()  # 隐藏输出目录（批量时自动保存到源文件夹）
            self.output_dir_hint.grid(row=5, column=0, columnspan=3, sticky=tk.W, padx=(10, 0), pady=2)  # 显示提示
            self.browse_button.config(command=self.browse_input_folder)
            self.detected_format_var.set("")  # 清空识别格式
            self.format_label.pack_forget()  # 隐藏识别格式
            self.batch_format_frame.pack(side=tk.LEFT)  # 显示批量格式选择
            self.input_path_var.set("")  # 清空路径
    
    def browse_input_file(self):
        """浏览选择输入文件"""
        file_path = filedialog.askopenfilename(
            title="选择输入文件",
            filetypes=[
                ("所有支持的格式", "*.md *.markdown *.docx *.html *.epub *.tex *.rst *.org *.odt *.rtf"),
                ("Markdown文件", "*.md *.markdown"),
                ("Word文档", "*.docx"),
                ("HTML文件", "*.html *.htm"),
                ("LaTeX文件", "*.tex *.ltx"),
                ("所有文件", "*.*")
            ]
        )
        if file_path:
            self.input_path_var.set(file_path)
    
    def browse_input_folder(self):
        """浏览选择输入文件夹（批量转换模式）"""
        dir_path = filedialog.askdirectory(title="选择输入文件夹（将扫描所有子文件夹）")
        if dir_path:
            self.input_path_var.set(dir_path)
    
    def browse_output_dir(self):
        """浏览选择输出目录"""
        dir_path = filedialog.askdirectory(title="选择输出目录")
        if dir_path:
            self.output_dir_var.set(dir_path)
    
    def on_input_file_changed(self, *args):
        """输入文件改变时的处理"""
        input_path = self.input_path_var.get()
        
        if not input_path or not os.path.exists(input_path):
            self.detected_format_var.set("未选择文件")
            return
        
        # 自动识别格式
        file_ext = Path(input_path).suffix.lower()
        detected_format = self.EXT_TO_FORMAT.get(file_ext, 'markdown')
        format_name = self.INPUT_FORMATS.get(detected_format, '未知格式')
        
        self.detected_format_var.set(f"{format_name} ({detected_format})")
        
        # 自动设置输出目录为输入文件所在目录
        if not self.output_dir_var.get():
            output_dir = os.path.dirname(input_path)
            self.output_dir_var.set(output_dir)
    
    def select_common_formats(self):
        """选择常用格式"""
        common_formats = ['docx', 'pdf', 'html5', 'epub', 'latex', 'plain']
        self.clear_all_formats()
        for fmt in common_formats:
            if fmt in self.output_format_vars:
                self.output_format_vars[fmt].set(True)
    
    def select_all_formats(self):
        """全选所有格式"""
        for var in self.output_format_vars.values():
            var.set(True)
    
    def clear_all_formats(self):
        """清除所有选择"""
        for var in self.output_format_vars.values():
            var.set(False)
    
    def show_advanced_options(self):
        """显示高级选项对话框"""
        dialog = tk.Toplevel(self.root)
        dialog.title("高级选项")
        dialog.geometry("750x420")  # 增加宽度和高度以容纳单选框
        dialog.resizable(False, False)
        dialog.transient(self.root)
        dialog.grab_set()
        
        # 使对话框居中
        dialog.update_idletasks()
        x = (dialog.winfo_screenwidth() // 2) - (dialog.winfo_width() // 2)
        y = (dialog.winfo_screenheight() // 2) - (dialog.winfo_height() // 2)
        dialog.geometry(f"+{x}+{y}")
        
        # 主框架
        main_frame = ttk.Frame(dialog, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # 标题
        title_label = ttk.Label(
            main_frame,
            text="⚙️ 高级选项设置",
            font=('Arial', 14, 'bold')
        )
        title_label.pack(anchor=tk.W, pady=(0, 20))
        
        # Pandoc选项区域
        pandoc_frame = ttk.LabelFrame(main_frame, text="Pandoc转换选项", padding="15")
        pandoc_frame.pack(fill=tk.X, pady=(0, 15))
        
        ttk.Checkbutton(
            pandoc_frame,
            text="生成独立文档 (--standalone)  适用于HTML、LaTeX等需要完整文档结构的格式",
            variable=self.standalone_var
        ).pack(anchor=tk.W, pady=5)
        
        ttk.Checkbutton(
            pandoc_frame,
            text="生成目录 (--toc)  在文档开头自动生成目录",
            variable=self.toc_var
        ).pack(anchor=tk.W, pady=5)
        
        ttk.Checkbutton(
            pandoc_frame,
            text="章节自动编号 (--number-sections)  为章节添加序号",
            variable=self.number_sections_var
        ).pack(anchor=tk.W, pady=5)
        
        # DOCX字体设置区域
        docx_frame = ttk.LabelFrame(main_frame, text="DOCX文档字体设置", padding="15")
        docx_frame.pack(fill=tk.X, pady=(0, 15))
        
        # 说明文字
        ttk.Label(
            docx_frame,
            text="💡 设置转换为DOCX格式时使用的默认字体和字号",
            font=('Arial', 9),
            foreground='gray'
        ).pack(anchor=tk.W, pady=(0, 15))
        
        # 字体选择 - 使用单选框（横向排列）
        font_label_frame = ttk.Frame(docx_frame)
        font_label_frame.pack(fill=tk.X, pady=(0, 10))
        ttk.Label(
            font_label_frame, 
            text="字体:", 
            font=('Arial', 9, 'bold')
        ).pack(side=tk.LEFT, padx=(0, 15))
        
        font_radio_frame = ttk.Frame(docx_frame)
        font_radio_frame.pack(fill=tk.X, pady=(0, 15))
        
        fonts = [
            ("Times New Roman", "Times New Roman"),
            ("Arial", "Arial"),
            ("Calibri", "Calibri"),
            ("默认（不设置）", "默认（不设置）")
        ]
        
        for text, value in fonts:
            ttk.Radiobutton(
                font_radio_frame,
                text=text,
                variable=self.docx_font_var,
                value=value
            ).pack(side=tk.LEFT, padx=(0, 20))
        
        # 字号选择 - 使用单选框（横向排列）
        fontsize_label_frame = ttk.Frame(docx_frame)
        fontsize_label_frame.pack(fill=tk.X, pady=(0, 10))
        ttk.Label(
            fontsize_label_frame, 
            text="字号:", 
            font=('Arial', 9, 'bold')
        ).pack(side=tk.LEFT, padx=(0, 15))
        
        fontsize_radio_frame = ttk.Frame(docx_frame)
        fontsize_radio_frame.pack(fill=tk.X, pady=(0, 5))
        
        fontsizes = [
            ("10pt", "10"),
            ("11pt", "11"),
            ("12pt", "12"),
            ("14pt", "14"),
            ("16pt", "16"),
            ("18pt", "18"),
            ("默认", "默认（不设置）")
        ]
        
        for text, value in fontsizes:
            ttk.Radiobutton(
                fontsize_radio_frame,
                text=text,
                variable=self.docx_fontsize_var,
                value=value
            ).pack(side=tk.LEFT, padx=(0, 15))
        
        # 按钮区域
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill=tk.X, pady=(10, 0))
        
        def on_ok():
            dialog.destroy()
        
        def on_reset():
            """重置为默认值"""
            self.standalone_var.set(True)
            self.toc_var.set(False)
            self.number_sections_var.set(False)
            self.docx_font_var.set("Times New Roman")
            self.docx_fontsize_var.set("12")
        
        ttk.Button(
            button_frame,
            text="确定",
            command=on_ok,
            style='Accent.TButton',
            width=12
        ).pack(side=tk.RIGHT, padx=(5, 0))
        
        ttk.Button(
            button_frame,
            text="恢复默认",
            command=on_reset,
            width=12
        ).pack(side=tk.RIGHT)
    
    def show_file_exists_dialog(self, filename):
        """显示文件已存在对话框，让用户选择处理方式"""
        # 如果已经记住了选择，直接返回
        if self.remember_choice and self.file_overwrite_policy:
            return self.file_overwrite_policy
        
        # 创建自定义对话框
        dialog = tk.Toplevel(self.root)
        dialog.title("文件已存在")
        dialog.geometry("600x400")  # 增加尺寸：从500x280改为600x400
        dialog.resizable(False, False)
        dialog.transient(self.root)
        dialog.grab_set()
        
        # 使对话框居中
        dialog.update_idletasks()
        x = (dialog.winfo_screenwidth() // 2) - (dialog.winfo_width() // 2)
        y = (dialog.winfo_screenheight() // 2) - (dialog.winfo_height() // 2)
        dialog.geometry(f"+{x}+{y}")
        
        # 用户选择结果
        user_choice = {'action': None}
        
        # 主框架
        main_frame = ttk.Frame(dialog, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # 警告图标和消息
        message_frame = ttk.Frame(main_frame)
        message_frame.pack(fill=tk.X, pady=(0, 20))
        
        # 图标
        icon_label = ttk.Label(
            message_frame,
            text="⚠️",
            font=('Arial', 32)
        )
        icon_label.pack(side=tk.LEFT, padx=(0, 15))
        
        # 消息文本
        text_frame = ttk.Frame(message_frame)
        text_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        ttk.Label(
            text_frame,
            text="文件已存在",
            font=('Arial', 12, 'bold')
        ).pack(anchor=tk.W)
        
        ttk.Label(
            text_frame,
            text=f"{os.path.basename(filename)}",
            font=('Arial', 10),
            foreground='blue'
        ).pack(anchor=tk.W, pady=(5, 0))
        
        ttk.Label(
            text_frame,
            text="请选择如何处理：",
            font=('Arial', 9)
        ).pack(anchor=tk.W, pady=(10, 0))
        
        # 选项框架
        options_frame = ttk.LabelFrame(main_frame, text="处理方式", padding="15")
        options_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 15))
        
        choice_var = tk.StringVar(value="overwrite")
        
        # 选项1：覆盖
        rb1 = ttk.Radiobutton(
            options_frame,
            text="覆盖现有文件 (将删除原文件)",
            variable=choice_var,
            value="overwrite"
        )
        rb1.pack(anchor=tk.W, pady=5)
        
        # 选项2：重命名
        rb2 = ttk.Radiobutton(
            options_frame,
            text="重命名新文件 (添加数字后缀，如 file_1.docx)",
            variable=choice_var,
            value="rename"
        )
        rb2.pack(anchor=tk.W, pady=5)
        
        # 选项3：跳过
        rb3 = ttk.Radiobutton(
            options_frame,
            text="跳过此文件 (不转换)",
            variable=choice_var,
            value="skip"
        )
        rb3.pack(anchor=tk.W, pady=5)
        
        # 记住选择复选框
        remember_var = tk.BooleanVar(value=False)
        remember_cb = ttk.Checkbutton(
            main_frame,
            text="记住我的选择，后续相同问题使用相同方式处理",
            variable=remember_var
        )
        remember_cb.pack(anchor=tk.W, pady=(0, 15))
        
        # 按钮框架
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill=tk.X)
        
        def on_confirm():
            user_choice['action'] = choice_var.get()
            if remember_var.get():
                self.file_overwrite_policy = choice_var.get()
                self.remember_choice = True
            dialog.destroy()
        
        def on_cancel():
            user_choice['action'] = 'skip'
            dialog.destroy()
        
        # 确定按钮
        confirm_button = ttk.Button(
            button_frame,
            text="确定",
            command=on_confirm,
            style='Accent.TButton'
        )
        confirm_button.pack(side=tk.RIGHT, padx=(5, 0))
        
        # 取消按钮
        cancel_button = ttk.Button(
            button_frame,
            text="取消",
            command=on_cancel
        )
        cancel_button.pack(side=tk.RIGHT)
        
        # 等待对话框关闭
        dialog.wait_window()
        
        return user_choice['action'] or 'skip'
    
    def get_unique_filename(self, filepath):
        """生成唯一的文件名（添加数字后缀）"""
        if not os.path.exists(filepath):
            return filepath
        
        directory = os.path.dirname(filepath)
        filename = os.path.basename(filepath)
        name, ext = os.path.splitext(filename)
        
        counter = 1
        while True:
            new_filename = f"{name}_{counter}{ext}"
            new_filepath = os.path.join(directory, new_filename)
            if not os.path.exists(new_filepath):
                return new_filepath
            counter += 1
    
    def _set_file_action(self, filepath):
        """在主线程中设置文件处理动作（供后台线程调用）"""
        action = self.show_file_exists_dialog(filepath)
        self._temp_file_action = action
    
    def remove_emoji_from_text(self, text):
        """移除或替换文本中的Emoji字符"""
        import re
        
        # 定义Emoji替换映射（替换为ASCII字符）
        emoji_map = {
            # 常用Emoji
            '⭐': '*',
            '✨': '*',
            '🎯': '[目标]',
            '📝': '[文档]',
            '🔧': '[工具]',
            '✅': '[√]',
            '❌': '[×]',
            '⚠️': '[!]',
            '💡': '[i]',
            '🚀': '>>',
            '🎉': '***',
            '📋': '[清单]',
            '🔄': '<->',
            '✓': '[√]',
            '✗': '[×]',
            # 时间相关
            '⏳': '[等待]',
            '⏰': '[时间]',
            '⌛': '[等待]',
            '⏱': '[计时]',  # 秒表
            '⏲': '[定时]',  # 定时器
            # 数字圈
            '⃣': '',  # 组合字符，直接移除
            # 变体选择符（不可见字符）
            '\uFE0F': '',  # Emoji变体选择符
            '\u20E3': '',  # 组合数字符号
            # 箱形图字符（树状图）
            '│': '|',
            '├': '|-',
            '└': '`-',
            '─': '-',
            '┌': '.-',
            '┐': '-.',
            '┘': '-\'',
            '┴': '-',
            '┬': '-',
            '┤': '-|',
            '┼': '+',
            # 方块字符（进度条等）
            '█': '#',
            '▓': '#',
            '▒': '=',
            '░': '-',
            '▀': '=',
            '▄': '=',
            '▌': '|',
            '▐': '|',
        }
        
        # 先替换常见的Emoji
        for emoji, replacement in emoji_map.items():
            text = text.replace(emoji, replacement)
        
        # 移除剩余的所有Emoji字符
        emoji_pattern = re.compile(
            "["
            "\U0001F600-\U0001F64F"  # 表情符号
            "\U0001F300-\U0001F5FF"  # 符号和象形文字
            "\U0001F680-\U0001F6FF"  # 交通和地图符号
            "\U0001F1E0-\U0001F1FF"  # 旗帜
            "\U00002600-\U000027BF"  # 杂项符号
            "\U0001F900-\U0001F9FF"  # 补充符号和象形文字
            "]+",
            flags=re.UNICODE
        )
        
        text = emoji_pattern.sub('', text)
        
        return text
    
    def get_docx_font_args(self):
        """获取DOCX字体设置参数"""
        args = []
        
        # 获取字体设置
        font = self.docx_font_var.get()
        fontsize = self.docx_fontsize_var.get()
        
        # 只有在不是"默认"时才添加参数
        if font and font != "默认（不设置）":
            # 使用metadata变量
            args.extend(['-M', f'mainfont={font}'])
        
        if fontsize and fontsize != "默认（不设置）":
            args.extend(['-M', f'fontsize={fontsize}'])
        
        return args
    
    def apply_docx_font(self, docx_file):
        """应用字体设置到生成的DOCX文件（后处理）"""
        try:
            from docx import Document
            from docx.shared import Pt
            
            font = self.docx_font_var.get()
            fontsize = self.docx_fontsize_var.get()
            
            # 如果都是默认，则不处理
            if (not font or font == "默认（不设置）") and \
               (not fontsize or fontsize == "默认（不设置）"):
                return True
            
            # 打开文档
            doc = Document(docx_file)
            
            # 设置默认字体
            if font and font != "默认（不设置）":
                # 设置样式的字体
                style = doc.styles['Normal']
                style.font.name = font
                # 设置中文字体
                style._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', font)
            
            if fontsize and fontsize != "默认（不设置）":
                style = doc.styles['Normal']
                style.font.size = Pt(int(fontsize))
            
            # 遍历所有段落应用字体
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if font and font != "默认（不设置）":
                        run.font.name = font
                        # 设置中文字体
                        run._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', font)
                    if fontsize and fontsize != "默认（不设置）":
                        run.font.size = Pt(int(fontsize))
            
            # 遍历所有表格
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if font and font != "默认（不设置）":
                                    run.font.name = font
                                    run._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', font)
                                if fontsize and fontsize != "默认（不设置）":
                                    run.font.size = Pt(int(fontsize))
            
            # 保存文档
            doc.save(docx_file)
            return True
            
        except ImportError:
            self.log_message("    ⚠ 警告: 未安装python-docx库，无法设置DOCX字体")
            self.log_message("    提示: pip install python-docx")
            return False
        except Exception as e:
            self.log_message(f"    ⚠ 设置DOCX字体失败: {str(e)}")
            return False
    
    def start_conversion(self):
        """开始转换"""
        # 检查Pandoc是否安装
        if not self.pandoc_installed:
            messagebox.showerror(
                "Pandoc 未安装",
                "Pandoc 未安装或未正确配置！\n\n"
                "请按以下步骤操作：\n"
                "1. 安装 Pandoc: https://pandoc.org/installing.html\n"
                "2. 点击'重新检测'按钮验证安装"
            )
            return
        
        input_path = self.input_path_var.get()
        mode = self.conversion_mode.get()
        
        # 验证输入
        if not input_path:
            if mode == "single":
                messagebox.showwarning("警告", "请选择输入文件")
            else:
                messagebox.showwarning("警告", "请选择输入文件夹")
            return
        
        if not os.path.exists(input_path):
            if mode == "single":
                messagebox.showerror("错误", "输入文件不存在")
            else:
                messagebox.showerror("错误", "输入文件夹不存在")
            return
        
        # 单文件模式需要检查输出目录
        if mode == "single":
            output_dir = self.output_dir_var.get()
            if not output_dir:
                messagebox.showwarning("警告", "请指定输出目录")
                return
            
            if not os.path.exists(output_dir):
                try:
                    os.makedirs(output_dir)
                except Exception as e:
                    messagebox.showerror("错误", f"无法创建输出目录:\n{str(e)}")
                    return
        
        # 检查是否至少选择了一个输出格式
        selected_formats = [fmt for fmt, var in self.output_format_vars.items() if var.get()]
        if not selected_formats:
            messagebox.showwarning("警告", "请至少选择一个输出格式")
            return
        
        # 清空日志
        self.log_text.delete(1.0, tk.END)
        
        # 重置文件覆盖策略（新的转换任务）
        self.file_overwrite_policy = None
        self.remember_choice = False
        
        # 根据模式启动不同的转换流程
        if mode == "single":
            output_dir = self.output_dir_var.get()
            self.status_var.set(f"正在转换到 {len(selected_formats)} 个格式...")
            self.progress.start()
            
            # 在新线程中执行单文件转换
            thread = threading.Thread(
                target=self.convert_files_thread,
                args=(input_path, output_dir, selected_formats)
            )
            thread.daemon = True
            thread.start()
        else:
            # 批量转换模式
            self.status_var.set("正在扫描文件夹...")
            self.progress.start()
            
            # 在新线程中执行批量转换
            thread = threading.Thread(
                target=self.batch_convert_thread,
                args=(input_path, selected_formats)
            )
            thread.daemon = True
            thread.start()
    
    def convert_files_thread(self, input_path, output_dir, selected_formats):
        """在后台线程中执行转换"""
        try:
            success_count = 0
            fail_count = 0
            
            input_file = Path(input_path)
            base_name = input_file.stem
            
            # 检测输入格式
            file_ext = input_file.suffix.lower()
            from_format = self.EXT_TO_FORMAT.get(file_ext, 'markdown')
            
            self.log_message(f"{'='*60}")
            self.log_message(f"输入文件: {input_path}")
            self.log_message(f"识别格式: {from_format}")
            self.log_message(f"输出目录: {output_dir}")
            self.log_message(f"选择格式数量: {len(selected_formats)}")
            self.log_message(f"{'='*60}\n")
            
            for to_format in selected_formats:
                try:
                    # 确定输出文件扩展名
                    if to_format == 'html5':
                        ext = 'html'
                    elif to_format == 'epub3':
                        ext = 'epub'
                    elif to_format == 'docbook5':
                        ext = 'xml'
                    elif to_format == 'plain':
                        ext = 'txt'
                    else:
                        ext = to_format
                    
                    output_file = os.path.join(output_dir, f"{base_name}.{ext}")
                    
                    # 检查文件是否存在
                    if os.path.exists(output_file):
                        # 在主线程中显示对话框
                        action = None
                        self.root.after(0, lambda: self._set_file_action(output_file))
                        # 等待用户选择
                        import time
                        while not hasattr(self, '_temp_file_action'):
                            time.sleep(0.1)
                        action = self._temp_file_action
                        delattr(self, '_temp_file_action')
                        
                        if action == 'skip':
                            self.log_message(f"    ⊘ 跳过 {to_format} (文件已存在)\n")
                            continue
                        elif action == 'rename':
                            output_file = self.get_unique_filename(output_file)
                            self.log_message(f"    📝 重命名为: {os.path.basename(output_file)}")
                        elif action == 'overwrite':
                            self.log_message(f"    ⚠️ 将覆盖现有文件")
                    
                    self.log_message(f"    ⏳ 正在转换到 {to_format}...")
                    
                    # 构建额外参数
                    extra_args = []
                    
                    if self.standalone_var.get():
                        extra_args.append('--standalone')
                    
                    if self.toc_var.get():
                        extra_args.append('--toc')
                    
                    if self.number_sections_var.get():
                        extra_args.append('--number-sections')
                    
                    # 执行转换
                    self.log_message(f"    ⏳ 正在转换到 {to_format}...")
                    
                    # PDF格式特殊处理
                    if to_format == 'pdf':
                        # 使用XeLaTeX引擎，支持Unicode和中文
                        pdf_args = extra_args.copy() if extra_args else []
                        pdf_args.extend([
                            '--pdf-engine=xelatex',
                            '-V', 'CJKmainfont=Microsoft YaHei',
                            '-V', 'mainfont=Microsoft YaHei',
                        ])
                        
                        # 如果输入是Markdown，先处理Emoji
                        if input_path.lower().endswith(('.md', '.markdown')):
                            self.log_message(f"    处理Emoji字符...")
                            # 读取文件
                            with open(input_path, 'r', encoding='utf-8') as f:
                                content = f.read()
                            
                            # 移除Emoji
                            content = self.remove_emoji_from_text(content)
                            
                            # 创建临时文件
                            import tempfile
                            temp_file = tempfile.NamedTemporaryFile(
                                mode='w',
                                encoding='utf-8',
                                suffix='.md',
                                delete=False
                            )
                            temp_file.write(content)
                            temp_file_path = temp_file.name
                            temp_file.close()
                            
                            try:
                                # 使用临时文件转换
                                output = pypandoc.convert_file(
                                    temp_file_path,
                                    to_format,
                                    outputfile=output_file,
                                    extra_args=pdf_args
                                )
                            finally:
                                # 删除临时文件
                                os.unlink(temp_file_path)
                        else:
                            # 非Markdown文件直接转换
                            output = pypandoc.convert_file(
                                input_path,
                                to_format,
                                format=from_format,
                                outputfile=output_file,
                                extra_args=pdf_args
                            )
                    else:
                        # 其他格式正常转换
                        output = pypandoc.convert_file(
                            input_path,
                            to_format,
                            format=from_format,
                            outputfile=output_file,
                            extra_args=extra_args
                        )
                    
                    # 如果是DOCX格式，应用字体设置
                    if to_format == 'docx':
                        self.log_message(f"    [工具] 正在应用DOCX字体设置...")
                        self.apply_docx_font(output_file)
                    
                    self.log_message(f"✓ 成功: {output_file}\n")
                    success_count += 1
                    
                except Exception as e:
                    self.log_message(f"✗ 失败 ({to_format}): {str(e)}\n")
                    fail_count += 1
            
            # 转换完成
            self.log_message(f"{'='*60}")
            self.log_message(f"转换完成!")
            self.log_message(f"成功: {success_count} 个, 失败: {fail_count} 个")
            self.log_message(f"{'='*60}")
            
            self.root.after(0, lambda: self.conversion_complete(success_count, fail_count, output_dir))
            
        except Exception as e:
            error_msg = str(e)
            self.root.after(0, lambda: self.conversion_error(error_msg))
    
    def batch_convert_thread(self, root_dir, selected_formats):
        """批量转换线程 - 扫描文件夹并转换所有支持的文件"""
        try:
            # 获取用户选择的输入格式
            input_format = self.batch_input_format.get()
            
            # 根据选择的输入格式确定要扫描的扩展名
            format_to_ext = {
                'md': ['.md', '.markdown'],
                'html': ['.html', '.htm'],
                'docx': ['.docx'],
                'txt': ['.txt'],
                'pptx': ['.pptx'],
                'pdf': ['.pdf']
            }
            
            target_extensions = format_to_ext.get(input_format, ['.md'])
            
            # 扫描所有文件
            self.log_message(f"{'='*60}")
            self.log_message(f"批量转换模式")
            self.log_message(f"扫描目录: {root_dir}")
            self.log_message(f"输入格式: {input_format.upper()}")
            self.log_message(f"目标格式: {', '.join(selected_formats)}")
            self.log_message(f"{'='*60}\n")
            
            # 使用Path.rglob递归扫描所有文件
            all_files = []
            root_path = Path(root_dir)
            
            self.log_message(f"正在扫描 {input_format.upper()} 格式的文件...")
            
            for ext in target_extensions:
                # 使用rglob递归查找所有匹配的文件
                matching_files = list(root_path.rglob(f"*{ext}"))
                all_files.extend(matching_files)
            
            # 去重（避免同一文件被多次添加）
            all_files = list(set(all_files))
            
            if not all_files:
                self.log_message(f"⚠ 警告: 在指定目录中未找到 {input_format.upper()} 格式的文件")
                self.root.after(0, lambda: self.conversion_complete(0, 0, root_dir))
                return
            
            self.log_message(f"✓ 找到 {len(all_files)} 个文件待转换\n")
            
            # 统计信息
            total_conversions = len(all_files) * len(selected_formats)
            success_count = 0
            fail_count = 0
            processed_files = 0
            
            # 逐个文件进行转换
            for input_file in all_files:
                processed_files += 1
                
                # 计算总体进度百分比
                overall_progress = (processed_files / len(all_files)) * 100
                
                # 更新状态
                self.root.after(0, lambda p=processed_files, t=len(all_files), pct=overall_progress: 
                    self.status_var.set(f"正在处理 {p}/{t} 个文件 ({pct:.1f}%)"))
                
                input_path = str(input_file)
                file_name = input_file.name
                output_dir = str(input_file.parent)  # 输出到源文件所在目录
                
                # 检测输入格式
                file_ext = input_file.suffix.lower()
                from_format = self.EXT_TO_FORMAT.get(file_ext, 'markdown')
                
                self.log_message(f"{'─'*60}")
                self.log_message(f"[{processed_files}/{len(all_files)}] ({overall_progress:.1f}%) 处理: {file_name}")
                self.log_message(f"    📂 路径: {input_path}")
                self.log_message(f"    📄 格式: {from_format}")
                
                # 转换到所有选定的格式
                format_count = 0
                for to_format in selected_formats:
                    format_count += 1
                    try:
                        # 确定输出文件扩展名
                        if to_format == 'html5':
                            ext = 'html'
                        elif to_format == 'epub3':
                            ext = 'epub'
                        elif to_format == 'docbook5':
                            ext = 'xml'
                        elif to_format == 'plain':
                            ext = 'txt'
                        else:
                            ext = to_format
                        
                        base_name = input_file.stem
                        output_file = os.path.join(output_dir, f"{base_name}.{ext}")
                        
                        # 如果输出文件与输入文件相同，跳过
                        if output_file == input_path:
                            self.log_message(f"    [{format_count}/{len(selected_formats)}] ⊘ 跳过 {to_format} (与源文件相同)")
                            continue
                        
                        # 检查文件是否存在
                        if os.path.exists(output_file):
                            # 在主线程中显示对话框
                            action = None
                            self.root.after(0, lambda f=output_file: self._set_file_action(f))
                            # 等待用户选择
                            import time
                            while not hasattr(self, '_temp_file_action'):
                                time.sleep(0.1)
                            action = self._temp_file_action
                            delattr(self, '_temp_file_action')
                            
                            if action == 'skip':
                                self.log_message(f"    [{format_count}/{len(selected_formats)}] ⊘ 跳过 {to_format} (文件已存在)")
                                continue
                            elif action == 'rename':
                                output_file = self.get_unique_filename(output_file)
                                self.log_message(f"    [{format_count}/{len(selected_formats)}] 📝 重命名为: {os.path.basename(output_file)}")
                            elif action == 'overwrite':
                                self.log_message(f"    [{format_count}/{len(selected_formats)}] ⚠️ 将覆盖现有文件")
                        
                        # 显示当前转换进度
                        format_progress = (format_count / len(selected_formats)) * 100
                        self.log_message(f"    [{format_count}/{len(selected_formats)}] ({format_progress:.0f}%) ⏳ 正在转换到 {to_format.upper()}...")
                        
                        # 构建额外参数
                        extra_args = []
                        
                        if self.standalone_var.get():
                            extra_args.append('--standalone')
                        
                        if self.toc_var.get():
                            extra_args.append('--toc')
                        
                        if self.number_sections_var.get():
                            extra_args.append('--number-sections')
                        
                        # 执行转换（PDF格式特殊处理）
                        if to_format == 'pdf':
                            # 使用XeLaTeX引擎，支持Unicode和中文
                            pdf_args = extra_args.copy() if extra_args else []
                            pdf_args.extend([
                                '--pdf-engine=xelatex',
                                '-V', 'CJKmainfont=Microsoft YaHei',
                                '-V', 'mainfont=Microsoft YaHei',
                            ])
                            
                            # 如果输入是Markdown，先处理Emoji
                            if str(input_path).lower().endswith(('.md', '.markdown')):
                                # 读取文件
                                with open(input_path, 'r', encoding='utf-8') as f:
                                    content = f.read()
                                
                                # 移除Emoji
                                content = self.remove_emoji_from_text(content)
                                
                                # 创建临时文件
                                import tempfile
                                temp_file = tempfile.NamedTemporaryFile(
                                    mode='w',
                                    encoding='utf-8',
                                    suffix='.md',
                                    delete=False
                                )
                                temp_file.write(content)
                                temp_file_path = temp_file.name
                                temp_file.close()
                                
                                try:
                                    # 使用临时文件转换
                                    output = pypandoc.convert_file(
                                        temp_file_path,
                                        to_format,
                                        outputfile=output_file,
                                        extra_args=pdf_args
                                    )
                                finally:
                                    # 删除临时文件
                                    os.unlink(temp_file_path)
                            else:
                                # 非Markdown文件直接转换
                                output = pypandoc.convert_file(
                                    input_path,
                                    to_format,
                                    format=from_format,
                                    outputfile=output_file,
                                    extra_args=pdf_args
                                )
                        else:
                            # 其他格式正常转换
                            output = pypandoc.convert_file(
                                input_path,
                                to_format,
                                format=from_format,
                                outputfile=output_file,
                                extra_args=extra_args
                            )
                        
                        # 如果是DOCX格式，应用字体设置
                        if to_format == 'docx':
                            font = self.docx_font_var.get()
                            fontsize = self.docx_fontsize_var.get()
                            if (font and font != "默认（不设置）") or (fontsize and fontsize != "默认（不设置）"):
                                self.log_message(f"        🔧 正在应用DOCX字体设置 (字体:{font}, 字号:{fontsize})...")
                                if self.apply_docx_font(output_file):
                                    self.log_message(f"        ✓ 字体设置完成")
                        
                        self.log_message(f"        ✓ 转换成功 → {os.path.basename(output_file)}")
                        success_count += 1
                        
                    except Exception as e:
                        self.log_message(f"    ✗ {to_format} 失败: {str(e)}")
                        fail_count += 1
                
                self.log_message("")  # 空行分隔
            
            # 转换完成
            self.log_message(f"{'='*60}")
            self.log_message(f"批量转换完成!")
            self.log_message(f"处理文件: {processed_files} 个")
            self.log_message(f"总转换数: {total_conversions} 次")
            self.log_message(f"成功: {success_count} 次, 失败: {fail_count} 次")
            self.log_message(f"{'='*60}")
            
            self.root.after(0, lambda: self.conversion_complete(success_count, fail_count, root_dir))
            
        except Exception as e:
            error_msg = str(e)
            self.root.after(0, lambda: self.conversion_error(error_msg))
    
    def log_message(self, message):
        """在日志区域添加消息"""
        self.root.after(0, lambda: self._add_log_message(message))
    
    def _add_log_message(self, message):
        """实际添加日志消息的方法"""
        self.log_text.insert(tk.END, message + "\n")
        self.log_text.see(tk.END)
        self.root.update()
    
    def conversion_complete(self, success_count, fail_count, output_dir):
        """转换完成后的处理"""
        self.progress.stop()
        self.status_var.set(f"转换完成! 成功: {success_count}, 失败: {fail_count}")
        
        msg = f"转换完成!\n\n成功: {success_count} 个文件\n失败: {fail_count} 个文件\n\n输出目录:\n{output_dir}"
        
        if success_count > 0:
            result = messagebox.askquestion(
                "转换完成",
                msg + "\n\n是否打开输出目录?",
                icon='info'
            )
            if result == 'yes':
                self.open_folder(output_dir)
        else:
            messagebox.showwarning("转换完成", msg)
    
    def conversion_error(self, error_message):
        """转换失败后的处理"""
        self.progress.stop()
        self.status_var.set("转换失败")
        messagebox.showerror("错误", f"转换过程中出现错误:\n{error_message}")
    
    def open_folder(self, path):
        """打开文件夹"""
        if sys.platform == 'win32':
            os.startfile(path)
        elif sys.platform == 'darwin':
            os.system(f'open "{path}"')
        else:
            os.system(f'xdg-open "{path}"')
    
    def recheck_environment(self):
        """重新检测环境"""
        # 重新检测Pandoc
        old_pandoc_status = self.pandoc_installed
        self.check_pandoc_availability()
        
        # 重新检测LaTeX
        old_latex_status = self.latex_installed
        self.check_latex_availability()
        
        # 更新Pandoc状态显示
        pandoc_status_text = f"Pandoc: {self.pandoc_version}"
        pandoc_status_color = 'green' if self.pandoc_installed else 'red'
        pandoc_status_icon = '✓' if self.pandoc_installed else '✗'
        self.pandoc_status_label.config(
            text=f"{pandoc_status_icon} {pandoc_status_text}",
            foreground=pandoc_status_color
        )
        
        # 更新LaTeX状态显示
        latex_status_text = "LaTeX: 已安装" if self.latex_installed else "LaTeX: 未安装"
        latex_status_color = 'green' if self.latex_installed else 'orange'
        latex_status_icon = '✓' if self.latex_installed else '✗'
        self.latex_status_label.config(
            text=f"{latex_status_icon} {latex_status_text}",
            foreground=latex_status_color
        )
        
        # 更新转换按钮状态
        if self.pandoc_installed:
            self.convert_button.config(state='normal')
            self.status_var.set("✓ 准备就绪 - 可以开始转换")
            self.status_label.config(foreground='green')
        else:
            self.convert_button.config(state='disabled')
            self.status_var.set("✗ Pandoc 未安装 - 请先安装 Pandoc")
            self.status_label.config(foreground='red')
        
        # 更新PDF选项状态（需要重新创建输出格式复选框，这里简化处理）
        # 显示检测结果消息
        changes = []
        if old_pandoc_status != self.pandoc_installed:
            if self.pandoc_installed:
                changes.append("✓ Pandoc 检测成功")
            else:
                changes.append("✗ Pandoc 未检测到")
        
        if old_latex_status != self.latex_installed:
            if self.latex_installed:
                changes.append("✓ LaTeX 检测成功")
            else:
                changes.append("✗ LaTeX 未检测到")
        
        if changes:
            messagebox.showinfo("环境检测", "\n".join(changes) + "\n\n" + 
                              ("提示: 如果刚安装了LaTeX，PDF选项需要重启程序才能生效。" if self.latex_installed and old_latex_status != self.latex_installed else ""))
        else:
            messagebox.showinfo("环境检测", 
                              f"Pandoc: {'已安装' if self.pandoc_installed else '未安装'}\n"
                              f"LaTeX: {'已安装' if self.latex_installed else '未安装'}")
    
    def show_help(self):
        """显示使用说明对话框"""
        help_window = tk.Toplevel(self.root)
        help_window.title(f"使用说明 - ChangoConverter 全功能文档转换器 v{self.VERSION}")
        help_window.geometry("700x600")
        help_window.resizable(True, True)
        
        # 创建主框架
        help_frame = ttk.Frame(help_window, padding="15")
        help_frame.pack(fill=tk.BOTH, expand=True)
        
        # 标题
        title_label = ttk.Label(
            help_frame,
            text="📖 ChangoConverter 全功能文档转换器 - 使用说明",
            font=('Arial', 13, 'bold')
        )
        title_label.pack(pady=(0, 15))
        
        # 创建滚动文本框
        text_frame = ttk.Frame(help_frame)
        text_frame.pack(fill=tk.BOTH, expand=True)
        
        scrollbar = ttk.Scrollbar(text_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        help_text = tk.Text(
            text_frame,
            wrap=tk.WORD,
            font=('Microsoft YaHei UI', 10),
            yscrollcommand=scrollbar.set,
            padx=10,
            pady=10
        )
        help_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.config(command=help_text.yview)
        
        # 使用说明内容
        help_content = f"""
软件版本: v{self.VERSION}
Pandoc 版本: {self.pandoc_version}

═══════════════════════════════════════════════════════

🎯 功能介绍

这是一个全功能文档格式转换工具，支持 30+ 种输入格式和 40+ 种输出格式的互相转换。

✨ 新功能: 批量转换模式
支持选择一个文件夹，自动扫描所有子文件夹中的文档文件并批量转换！

═══════════════════════════════════════════════════════

📝 使用步骤

🔹 单文件转换模式

1️⃣ 选择转换模式
   选择"单文件转换"单选按钮

2️⃣ 选择输入文件
   点击"浏览"按钮选择要转换的文档文件
   ✓ 支持的格式：Markdown、Word、HTML、LaTeX、EPUB、PDF 等
   ✓ 程序会自动识别文件格式

3️⃣ 设置输出目录
   默认为输入文件所在目录
   ✓ 可以点击"浏览"按钮选择其他目录
   ✓ 输出文件名自动根据输入文件生成

🔹 批量转换模式（新功能）

1️⃣ 选择转换模式
   选择"批量转换（文件夹）"单选按钮

2️⃣ 选择输入文件夹
   点击"浏览"按钮选择包含文档的文件夹
   ✓ 会自动递归扫描所有子文件夹
   ✓ 自动识别所有支持的文件格式
   ✓ 转换后的文件保存到源文件所在的文件夹

🔸 通用步骤（单文件和批量模式都适用）

1️⃣ 选择输出格式
   在格式列表中勾选需要输出的格式
   ✓ 可以同时选择多个格式
   ✓ 常用格式在顶部优先显示
   ✓ 快速选择按钮：
      • 常用格式 - 选择最常用的 6 种格式
      • 全选 - 选择所有格式
      • 清除 - 取消所有选择

2️⃣ 高级选项（可选）
   ☑ 独立文档 (--standalone)
      生成完整的独立文档，包含文档头和尾
   
   ☑ 生成目录 (--toc)
      自动生成文档目录
   
   ☑ 章节编号 (--number-sections)
      为章节标题添加自动编号

3️⃣ 开始转换
   点击"开始转换"按钮
   ✓ 单文件模式：转换到指定输出目录
   ✓ 批量模式：扫描并转换所有文件
   ✓ 转换日志实时显示进度
   ✓ 完成后可直接打开输出目录

═══════════════════════════════════════════════════════

📌 常用格式说明

• Markdown (.md)
  轻量级标记语言，适合写作和笔记

• HTML5 (.html)
  网页格式，可在浏览器中查看

• DOCX (.docx)
  Microsoft Word 文档，最通用的文档格式

• PDF (.pdf)
  便携式文档格式，适合打印和正式分享
  ⚠️ 需要先安装 LaTeX（见下方安装指南）
  推荐：Windows 用户安装 MiKTeX
  首次转换可能较慢（自动下载必要组件）

• Plain Text (.txt)
  纯文本格式，去除所有格式

• PowerPoint (.pptx)
  演示文稿格式

═══════════════════════════════════════════════════════

🔧 支持的格式类型

📥 输入格式（30+）
  Markdown（多种变体）、Word、HTML、LaTeX、EPUB、
  reStructuredText、Org-mode、MediaWiki、DocBook、
  Jupyter Notebook、RTF、ODT 等

📤 输出格式（40+）
  除输入格式外，还支持：
  • 幻灯片：Beamer、reveal.js、Slidy、S5 等
  • 电子书：EPUB2、EPUB3
  • 排版：ConTeXt、Texinfo、InDesign ICML
  • 其他：AsciiDoc、Textile、ZimWiki 等

═══════════════════════════════════════════════════════

💡 使用技巧

1. 批量转换的高效方式
   • 使用批量转换模式处理整个文件夹
   • 自动递归扫描所有子文件夹
   • 一次性转换所有支持的文件格式
   • 勾选多个输出格式，一次完成所有转换

2. 格式转换建议
   • Markdown → DOCX：适合提交文档
   • Markdown → PDF：适合打印和分享
   • DOCX → Markdown：便于版本控制
   • HTML → PDF：网页转文档

3. PDF 转换说明
   PDF 格式需要先安装 LaTeX 引擎，请参考下方的安装指南

4. 文档质量
   使用"独立文档"选项可以获得更完整的输出
   使用"生成目录"可以让长文档更易阅读

═══════════════════════════════════════════════════════

📦 LaTeX 安装指南（PDF 转换必需）

PDF 格式转换需要安装 LaTeX 引擎，推荐以下方案：

▶ Windows 系统 - MiKTeX（推荐）

1. 下载 MiKTeX
   • 官网：https://miktex.org/download
   • 选择"Download MiKTeX"
   • 推荐下载：Basic MiKTeX Installer（约 200MB）

2. 安装步骤
   ① 运行下载的安装程序
   ② 选择安装类型：
      - "Install just for me"（仅当前用户）- 推荐
      - "Install for all users"（所有用户）- 需要管理员权限
   ③ 选择安装路径（默认即可）
   ④ 设置："Install missing packages on-the-fly"（自动安装缺失包）
      选择"Yes"或"Ask me first"
   ⑤ 点击"Start"开始安装
   ⑥ 安装完成后，MiKTeX 会自动配置环境变量

3. 验证安装
   • 打开命令提示符（CMD）
   • 输入：pdflatex --version
   • 如果显示版本信息，说明安装成功

4. 首次使用
   • 首次转换 PDF 时会自动下载必要的包
   • 可能需要几分钟时间
   • 请保持网络连接

▶ Windows 系统 - TeX Live（完整版）

如果需要完整的 LaTeX 功能：
   • 官网：https://tug.org/texlive/
   • 下载：install-tl-windows.exe
   • 注意：完整安装约 4-7GB
   • 安装时间较长（30-60分钟）

▶ macOS 系统 - MacTeX

1. 下载 MacTeX
   • 官网：https://www.tug.org/mactex/
   • 下载 MacTeX.pkg（约 4GB）

2. 安装步骤
   • 双击 .pkg 文件
   • 按照向导完成安装
   • 安装时间约 30-60 分钟

3. 验证安装
   • 打开终端（Terminal）
   • 输入：pdflatex --version

▶ Linux 系统 - TeX Live

Ubuntu/Debian:
   sudo apt-get update
   sudo apt-get install texlive-full

Fedora/CentOS:
   sudo dnf install texlive-scheme-full

Arch Linux:
   sudo pacman -S texlive-most

验证安装:
   pdflatex --version

▶ 轻量级选择 - TinyTeX（跨平台）

适合只需要 PDF 转换的用户：
   • 官网：https://yihui.org/tinytex/
   • 大小：约 100MB
   • 支持 Windows、macOS、Linux

Windows 安装：
   在 PowerShell 中运行：
   wget -qO- "https://yihui.org/tinytex/install-bin-windows.bat" | cmd

═══════════════════════════════════════════════════════

❓ 常见问题

Q: 转换 PDF 失败，显示"pdflatex not found"？
A: 这说明未安装 LaTeX。请按照上面的指南安装 MiKTeX 或其他 LaTeX 发行版
   安装后需要：
   • 重启命令提示符或 PowerShell
   • 或重启本程序
   • 确保 LaTeX 在系统 PATH 中

Q: PDF 转换很慢或卡住？
A: 首次转换 PDF 时，MiKTeX 会自动下载必要的包，可能需要几分钟
   请耐心等待，后续转换会快很多
   确保有稳定的网络连接

Q: 转换失败怎么办？
A: 检查转换日志中的错误信息，可能是：
   • 输入文件格式不正确
   • 缺少必要的依赖（如 LaTeX）
   • 文件包含不支持的特殊元素
   • LaTeX 包缺失（MiKTeX 会自动下载）

Q: 如何提高转换质量？
A: • 使用标准的 Markdown 语法
   • 勾选"独立文档"选项
   • 确保输入文件格式正确
   • 对于 PDF：确保 LaTeX 正确安装

Q: 支持哪些 Markdown 语法？
A: 支持 Pandoc 扩展的 Markdown，包括：
   • 表格、脚注、数学公式
   • 代码块语法高亮
   • 元数据块（YAML）
   • 上标、下标、删除线

Q: MiKTeX 和 TeX Live 选哪个？
A: • MiKTeX - 推荐 Windows 用户，轻量、易用
   • TeX Live - 跨平台，功能完整
   • TinyTeX - 最小安装，仅 PDF 转换
   建议：普通用户选 MiKTeX，专业用户选 TeX Live

═══════════════════════════════════════════════════════

🔗 相关链接

Pandoc:
• Pandoc 官网: https://pandoc.org/
• Pandoc 文档: https://pandoc.org/MANUAL.html
• Pandoc 安装: https://pandoc.org/installing.html

LaTeX 发行版（PDF 转换必需）:
• MiKTeX (Windows 推荐): https://miktex.org/download
• TeX Live (跨平台): https://tug.org/texlive/
• MacTeX (macOS): https://www.tug.org/mactex/
• TinyTeX (轻量级): https://yihui.org/tinytex/

═══════════════════════════════════════════════════════

© 2024 ChangoConverter 全功能文档转换器 v{self.VERSION}
基于 Pandoc {self.pandoc_version} 构建
        """
        
        help_text.insert(1.0, help_content)
        help_text.config(state=tk.DISABLED)  # 只读
        
        # 关闭按钮
        close_button = ttk.Button(
            help_frame,
            text="关闭",
            command=help_window.destroy,
            style='Accent.TButton'
        )
        close_button.pack(pady=(15, 0))
        
        # 居中显示窗口
        help_window.transient(self.root)
        help_window.grab_set()
        
        # 计算居中位置
        help_window.update_idletasks()
        x = (help_window.winfo_screenwidth() // 2) - (help_window.winfo_width() // 2)
        y = (help_window.winfo_screenheight() // 2) - (help_window.winfo_height() // 2)
        help_window.geometry(f"+{x}+{y}")

def main():
    root = tk.Tk()
    app = ChangoConverter(root)
    root.mainloop()

if __name__ == "__main__":
    main()

